"""
Document generation service for SDA evaluations.
Generates Word documents using python-docx.
"""
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..models.evaluation import Evaluation, ConclusionType
from ..models.service import Service
from .conclusion import get_conclusion_description

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
OUTPUT_DIR = Path(__file__).parent.parent.parent / "output"


def generate_evaluation_document(
    evaluation: Evaluation,
    service: Service,
    threats: list,
    safeguards: list
) -> str:
    """
    Generate a Word document for the SDA evaluation.

    Returns the path to the generated document.
    """
    # Ensure output directory exists
    OUTPUT_DIR.mkdir(exist_ok=True)

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header
    _add_header(doc, evaluation)

    # Title
    title = doc.add_heading('ANÁLISIS DE AMENAZAS - MEDIDAS DE SALVAGUARDA', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('EVALUACIÓN DE SERVICIO DISTINTO DE AUDITORÍA (SDA)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Section 1: General Information
    _add_section_general_info(doc, evaluation, service)

    # Section 2: Legal Gate Result
    _add_section_legal_gate(doc, evaluation)

    # Section 3: Threats Analysis
    _add_section_threats(doc, threats)

    # Section 4: Safeguards
    _add_section_safeguards(doc, safeguards)

    # Section 5: Conclusion
    _add_section_conclusion(doc, evaluation)

    # Section 6: Signatures
    _add_section_signatures(doc, evaluation)

    # Save document
    filename = f"SDA_{evaluation.reference_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))

    return str(output_path)


def _add_header(doc: Document, evaluation: Evaluation):
    """Add document header with Mazars branding"""
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add Mazars text (in production, would add logo image)
    run = header_para.add_run("FORVIS MAZARS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Reference number
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref_run = ref_para.add_run(f"Ref: {evaluation.reference_number}")
    ref_run.font.size = Pt(10)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"Fecha: {evaluation.created_at.strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(10)

    doc.add_paragraph()


def _add_section_general_info(doc: Document, evaluation: Evaluation, service: Service):
    """Add general information section"""
    doc.add_heading('1. INFORMACIÓN GENERAL', level=2)

    # Create table for general info
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    rows_data = [
        ('Entidad Auditada:', evaluation.entity_name or 'No especificada'),
        ('Tipo de Entidad:', 'Entidad de Interés Público (EIP)' if evaluation.entity_type.value == 'EIP' else 'Entidad NO EIP'),
        ('Tipo de Relación:', _get_relation_type_text(evaluation.relation_type.value)),
        ('Servicio Solicitado:', f"{service.code} - {service.name}" if service else 'No especificado'),
        ('Auditor Responsable:', evaluation.auditor_name or 'No especificado'),
    ]

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Bold the label
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()


def _add_section_legal_gate(doc: Document, evaluation: Evaluation):
    """Add legal gate result section"""
    doc.add_heading('2. RESULTADO DEL GATE LEGAL', level=2)

    passed = evaluation.legal_gate_passed == 1
    status = "✓ APROBADO" if passed else "✗ NO APROBADO"

    para = doc.add_paragraph()
    run = para.add_run(f"Estado: {status}")
    run.bold = True
    if passed:
        run.font.color.rgb = RGBColor(0, 128, 0)
    else:
        run.font.color.rgb = RGBColor(255, 0, 0)

    if evaluation.legal_gate_reason:
        doc.add_paragraph(f"Motivo: {evaluation.legal_gate_reason}")

    doc.add_paragraph()


def _add_section_threats(doc: Document, threats: list):
    """Add threats analysis section"""
    doc.add_heading('3. ANÁLISIS DE AMENAZAS', level=2)

    if not threats:
        doc.add_paragraph("No se han identificado amenazas significativas para este servicio.")
    else:
        # Create table for threats
        table = doc.add_table(rows=len(threats) + 1, cols=3)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Tipo de Amenaza'
        header_cells[1].text = 'Significatividad'
        header_cells[2].text = 'Observaciones'

        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for i, threat in enumerate(threats, 1):
            row = table.rows[i].cells
            row[0].text = threat.get('name', '')
            row[1].text = _get_significance_text(threat.get('significance', 'MEDIUM'))
            row[2].text = threat.get('notes', '') or ''

    doc.add_paragraph()


def _add_section_safeguards(doc: Document, safeguards: list):
    """Add safeguards section"""
    doc.add_heading('4. MEDIDAS DE SALVAGUARDA APLICADAS', level=2)

    if not safeguards:
        doc.add_paragraph("No se han seleccionado medidas de salvaguarda.")
    else:
        for i, safeguard in enumerate(safeguards, 1):
            para = doc.add_paragraph(style='List Number')
            para.add_run(safeguard.get('description', ''))
            if safeguard.get('notes'):
                notes_para = doc.add_paragraph()
                notes_para.paragraph_format.left_indent = Inches(0.5)
                notes_run = notes_para.add_run(f"Nota: {safeguard['notes']}")
                notes_run.italic = True

    doc.add_paragraph()


def _add_section_conclusion(doc: Document, evaluation: Evaluation):
    """Add conclusion section"""
    doc.add_heading('5. CONCLUSIÓN', level=2)

    if evaluation.conclusion:
        conclusion_info = get_conclusion_description(evaluation.conclusion)

        # Conclusion code and title
        para = doc.add_paragraph()
        code_run = para.add_run(f"{evaluation.conclusion.value}: {conclusion_info['title']}")
        code_run.bold = True
        code_run.font.size = Pt(14)

        # Set color based on conclusion
        color_map = {
            'green': RGBColor(0, 128, 0),
            'blue': RGBColor(0, 0, 255),
            'yellow': RGBColor(255, 165, 0),
            'orange': RGBColor(255, 140, 0),
            'red': RGBColor(255, 0, 0),
            'gray': RGBColor(128, 128, 128),
        }
        code_run.font.color.rgb = color_map.get(conclusion_info['color'], RGBColor(0, 0, 0))

        # Description
        doc.add_paragraph(conclusion_info['description'])

        # Additional notes
        if evaluation.conclusion_notes:
            notes_para = doc.add_paragraph()
            notes_para.add_run("Observaciones adicionales: ").bold = True
            notes_para.add_run(evaluation.conclusion_notes)

    doc.add_paragraph()


def _add_section_signatures(doc: Document, evaluation: Evaluation):
    """Add signatures section"""
    doc.add_heading('6. FIRMAS', level=2)

    # Create signature table
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Auditor signature
    table.rows[0].cells[0].text = "Auditor Responsable:"
    table.rows[0].cells[1].text = evaluation.auditor_name or "_____________________"

    # Date
    table.rows[1].cells[0].text = "Fecha:"
    table.rows[1].cells[1].text = "_____________________"

    # Signature line
    table.rows[2].cells[0].text = "Firma:"
    table.rows[2].cells[1].text = "_____________________"

    doc.add_paragraph()
    doc.add_paragraph()

    # Ethics Partner signature (if needed)
    if evaluation.conclusion in [ConclusionType.C3, ConclusionType.C4]:
        ethics_para = doc.add_paragraph()
        ethics_para.add_run("APROBACIÓN DEL SOCIO DE ÉTICA").bold = True

        ethics_table = doc.add_table(rows=3, cols=2)
        ethics_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        ethics_table.rows[0].cells[0].text = "Socio de Ética:"
        ethics_table.rows[0].cells[1].text = "_____________________"

        ethics_table.rows[1].cells[0].text = "Fecha:"
        ethics_table.rows[1].cells[1].text = "_____________________"

        ethics_table.rows[2].cells[0].text = "Firma:"
        ethics_table.rows[2].cells[1].text = "_____________________"


def _get_relation_type_text(relation_type: str) -> str:
    """Get Spanish text for relation type"""
    texts = {
        'AUDITADA': 'Entidad Auditada',
        'CADENA': 'Cadena de Control',
        'VINCULADA': 'Entidad Vinculada Significativa',
    }
    return texts.get(relation_type, relation_type)


def _get_significance_text(significance: str) -> str:
    """Get Spanish text for significance level"""
    texts = {
        'LOW': 'Baja',
        'MEDIUM': 'Media',
        'HIGH': 'Alta',
    }
    return texts.get(significance, significance)
